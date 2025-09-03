"""
Advanced File Format Parsers for Research Processing.

This module provides parsers for advanced file formats including PDF and DOCX files,
with robust error handling and content extraction capabilities.
"""

import re
from pathlib import Path
from typing import Any

try:
    import PyPDF2

    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    import docx
    from docx.document import Document

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import openpyxl

    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

from .research_processor import FileParser, Insight, ResearchFile


class PDFParser(FileParser):
    """Parser for PDF documents."""

    def __init__(self):
        super().__init__()
        self.supported_extensions = {".pdf"}
        self.supported_mime_types = {"application/pdf"}

    def can_parse(self, file_path: Path) -> bool:
        """Check if this parser can handle the file."""
        if not HAS_PYPDF2:
            self.logger.warning("PyPDF2 not available - cannot parse PDF files")
            return False
        return super().can_parse(file_path)

    async def parse_file(self, file_path: Path) -> ResearchFile:
        """Parse PDF file and extract content."""
        if not HAS_PYPDF2:
            return ResearchFile(
                path=file_path,
                content="",
                file_type="pdf",
                metadata={},
                processing_errors=["PyPDF2 not available"],
                confidence_score=0.0,
            )

        try:
            metadata = self._get_file_metadata(file_path)

            # Extract text from PDF
            text_content = await self._extract_pdf_text(file_path)

            # Extract insights from PDF structure
            insights = self._extract_pdf_insights(text_content, str(file_path))

            return ResearchFile(
                path=file_path,
                content=text_content,
                file_type="pdf",
                metadata={
                    "original_format": "pdf",
                    "pages_count": len(text_content.split("\f"))
                    if "\f" in text_content
                    else 1,
                    "estimated_word_count": len(text_content.split()),
                    **metadata.__dict__,
                },
                extracted_insights=[insight.content for insight in insights],
                confidence_score=0.7 if text_content.strip() else 0.1,
            )

        except Exception as e:
            self.logger.error(f"Failed to parse PDF file {file_path}: {e}")
            return ResearchFile(
                path=file_path,
                content="",
                file_type="pdf",
                metadata={},
                processing_errors=[str(e)],
                confidence_score=0.0,
            )

    async def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text content from PDF file."""
        text_content = ""

        try:
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)

                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content += page_text + "\n\n"

                        # Limit extraction to prevent memory issues
                        if page_num > 100:  # Max 100 pages
                            self.logger.warning(
                                "PDF has many pages, limiting to first 100"
                            )
                            break

                    except Exception as e:
                        self.logger.warning(
                            f"Failed to extract text from page {page_num}: {e}"
                        )
                        continue

        except Exception as e:
            self.logger.error(f"Failed to read PDF file: {e}")
            raise

        # Clean up the extracted text
        text_content = self._clean_pdf_text(text_content)

        return text_content

    def _clean_pdf_text(self, text: str) -> str:
        """Clean and normalize PDF extracted text."""
        # Remove excessive whitespace
        text = re.sub(r"\s+", " ", text)

        # Remove page markers
        text = re.sub(r"\f", "\n\n", text)

        # Fix common PDF extraction issues
        text = re.sub(r"(\w)-\s*(\w)", r"\1\2", text)  # Join hyphenated words
        text = re.sub(r"\n\s*\n\s*\n", "\n\n", text)  # Normalize line breaks

        # Remove headers/footers that repeat
        lines = text.split("\n")
        cleaned_lines = []

        for line in lines:
            line = line.strip()
            if len(line) > 3 and not self._is_likely_header_footer(line):
                cleaned_lines.append(line)

        return "\n".join(cleaned_lines)

    def _is_likely_header_footer(self, line: str) -> bool:
        """Check if a line is likely a header or footer."""
        # Common patterns for headers/footers
        patterns = [
            r"^\d+$",  # Just a page number
            r"^Page \d+",  # Page N
            r"^\d+\s*of\s*\d+$",  # N of M
            r"^Chapter \d+",  # Chapter N
            r"^©",  # Copyright
        ]

        for pattern in patterns:
            if re.match(pattern, line, re.IGNORECASE):
                return True

        # Very short lines are likely headers/footers
        if len(line) < 5:
            return True

        return False

    def _extract_pdf_insights(self, content: str, source_file: str) -> list[Insight]:
        """Extract insights from PDF content."""
        insights = []

        # Look for technical terms
        technical_terms = re.findall(
            r"\b(?:API|SDK|JSON|XML|HTTP|HTTPS|REST|GraphQL|SQL|NoSQL|Docker|Kubernetes|React|Vue|Angular|Node\.js|Python|Java|JavaScript|TypeScript|Git|CI/CD|DevOps|AWS|Azure|GCP)\b",
            content,
            re.IGNORECASE,
        )

        if technical_terms:
            unique_terms = list(set(term.lower() for term in technical_terms))
            insights.append(
                Insight(
                    content=f"Technical concepts found: {', '.join(unique_terms[:10])}",
                    source_file=source_file,
                    category="technology",
                    technical_concepts=unique_terms,
                    confidence_score=0.6,
                    importance_score=0.5,
                )
            )

        # Look for section headers
        headers = re.findall(r"^[A-Z][A-Z\s]{10,50}$", content, re.MULTILINE)
        if headers:
            insights.append(
                Insight(
                    content=f"Document sections: {', '.join(headers[:5])}",
                    source_file=source_file,
                    category="topic",
                    confidence_score=0.7,
                    importance_score=0.6,
                )
            )

        return insights


class DOCXParser(FileParser):
    """Parser for Microsoft Word documents."""

    def __init__(self):
        super().__init__()
        self.supported_extensions = {".docx", ".doc"}
        self.supported_mime_types = {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        }

    def can_parse(self, file_path: Path) -> bool:
        """Check if this parser can handle the file."""
        if not HAS_DOCX:
            self.logger.warning("python-docx not available - cannot parse DOCX files")
            return False
        # Only support .docx files, not .doc
        if file_path.suffix.lower() == ".doc":
            return False
        return super().can_parse(file_path)

    async def parse_file(self, file_path: Path) -> ResearchFile:
        """Parse DOCX file and extract content."""
        if not HAS_DOCX:
            return ResearchFile(
                path=file_path,
                content="",
                file_type="docx",
                metadata={},
                processing_errors=["python-docx not available"],
                confidence_score=0.0,
            )

        try:
            metadata = self._get_file_metadata(file_path)

            # Extract text from DOCX
            text_content, structure_info = await self._extract_docx_content(file_path)

            # Extract insights
            insights = self._extract_docx_insights(
                text_content, structure_info, str(file_path)
            )

            return ResearchFile(
                path=file_path,
                content=text_content,
                file_type="docx",
                metadata={
                    "original_format": "docx",
                    "paragraphs_count": structure_info.get("paragraphs_count", 0),
                    "tables_count": structure_info.get("tables_count", 0),
                    "headers_count": structure_info.get("headers_count", 0),
                    **metadata.__dict__,
                },
                extracted_insights=[insight.content for insight in insights],
                confidence_score=0.8 if text_content.strip() else 0.1,
            )

        except Exception as e:
            self.logger.error(f"Failed to parse DOCX file {file_path}: {e}")
            return ResearchFile(
                path=file_path,
                content="",
                file_type="docx",
                metadata={},
                processing_errors=[str(e)],
                confidence_score=0.0,
            )

    async def _extract_docx_content(
        self, file_path: Path
    ) -> tuple[str, dict[str, Any]]:
        """Extract text and structure information from DOCX file."""
        try:
            doc = docx.Document(str(file_path))

            text_parts = []
            structure_info = {
                "paragraphs_count": 0,
                "tables_count": 0,
                "headers_count": 0,
                "headers": [],
            }

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
                    structure_info["paragraphs_count"] += 1

                    # Check if it's a header (simplified heuristic)
                    if self._is_likely_header(paragraph, text):
                        structure_info["headers_count"] += 1
                        structure_info["headers"].append(text)

            # Extract tables
            for table in doc.tables:
                structure_info["tables_count"] += 1
                table_text = []

                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)

                    if row_text:
                        table_text.append(" | ".join(row_text))

                if table_text:
                    text_parts.append("Table:\n" + "\n".join(table_text))

            full_text = "\n\n".join(text_parts)

            return full_text, structure_info

        except Exception as e:
            self.logger.error(f"Failed to extract content from DOCX: {e}")
            raise

    def _is_likely_header(self, paragraph, text: str) -> bool:
        """Check if a paragraph is likely a header."""
        # Check if paragraph has header style
        if hasattr(paragraph, "style") and paragraph.style:
            style_name = str(paragraph.style.name).lower()
            if "heading" in style_name or "title" in style_name:
                return True

        # Check text characteristics
        if len(text) < 100 and text.isupper():
            return True

        if len(text) < 80 and not text.endswith("."):
            return True

        return False

    def _extract_docx_insights(
        self, content: str, structure_info: dict, source_file: str
    ) -> list[Insight]:
        """Extract insights from DOCX content and structure."""
        insights = []

        # Structure insights
        if structure_info["headers"]:
            insights.append(
                Insight(
                    content=f"Document structure: {', '.join(structure_info['headers'][:5])}",
                    source_file=source_file,
                    category="topic",
                    confidence_score=0.8,
                    importance_score=0.7,
                )
            )

        # Technical content
        technical_terms = re.findall(
            r"\b(?:API|SDK|JSON|XML|HTTP|HTTPS|REST|GraphQL|SQL|NoSQL|Docker|Kubernetes|React|Vue|Angular|Node\.js|Python|Java|JavaScript|TypeScript)\b",
            content,
            re.IGNORECASE,
        )

        if technical_terms:
            unique_terms = list(set(term.lower() for term in technical_terms))
            insights.append(
                Insight(
                    content=f"Technical topics: {', '.join(unique_terms[:8])}",
                    source_file=source_file,
                    category="technology",
                    technical_concepts=unique_terms,
                    confidence_score=0.7,
                    importance_score=0.6,
                )
            )

        return insights


class ExcelParser(FileParser):
    """Parser for Excel spreadsheets."""

    def __init__(self):
        super().__init__()
        self.supported_extensions = {".xlsx", ".xlsm"}
        self.supported_mime_types = {
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        }

    def can_parse(self, file_path: Path) -> bool:
        """Check if this parser can handle the file."""
        if not HAS_OPENPYXL:
            self.logger.warning("openpyxl not available - cannot parse Excel files")
            return False
        return super().can_parse(file_path)

    async def parse_file(self, file_path: Path) -> ResearchFile:
        """Parse Excel file and extract content."""
        if not HAS_OPENPYXL:
            return ResearchFile(
                path=file_path,
                content="",
                file_type="excel",
                metadata={},
                processing_errors=["openpyxl not available"],
                confidence_score=0.0,
            )

        try:
            metadata = self._get_file_metadata(file_path)

            # Extract data from Excel
            text_content, structure_info = await self._extract_excel_content(file_path)

            # Extract insights
            insights = self._extract_excel_insights(
                text_content, structure_info, str(file_path)
            )

            return ResearchFile(
                path=file_path,
                content=text_content,
                file_type="excel",
                metadata={
                    "original_format": "excel",
                    "sheets_count": structure_info.get("sheets_count", 0),
                    "total_rows": structure_info.get("total_rows", 0),
                    "sheet_names": structure_info.get("sheet_names", []),
                    **metadata.__dict__,
                },
                extracted_insights=[insight.content for insight in insights],
                confidence_score=0.6 if text_content.strip() else 0.1,
            )

        except Exception as e:
            self.logger.error(f"Failed to parse Excel file {file_path}: {e}")
            return ResearchFile(
                path=file_path,
                content="",
                file_type="excel",
                metadata={},
                processing_errors=[str(e)],
                confidence_score=0.0,
            )

    async def _extract_excel_content(
        self, file_path: Path
    ) -> tuple[str, dict[str, Any]]:
        """Extract data and structure from Excel file."""
        try:
            workbook = openpyxl.load_workbook(str(file_path), data_only=True)

            text_parts = []
            structure_info = {
                "sheets_count": len(workbook.worksheets),
                "sheet_names": [ws.title for ws in workbook.worksheets],
                "total_rows": 0,
            }

            for worksheet in workbook.worksheets:
                sheet_text = [f"Sheet: {worksheet.title}"]
                rows_processed = 0

                # Extract data from non-empty cells
                for row in worksheet.iter_rows(
                    max_row=min(100, worksheet.max_row)
                ):  # Limit rows
                    row_data = []
                    for cell in row:
                        if cell.value is not None:
                            row_data.append(str(cell.value))

                    if row_data:
                        sheet_text.append(" | ".join(row_data))
                        rows_processed += 1

                    # Limit processing to prevent memory issues
                    if rows_processed >= 50:
                        break

                if len(sheet_text) > 1:  # More than just the sheet title
                    text_parts.append("\n".join(sheet_text))
                    structure_info["total_rows"] += rows_processed

            full_text = "\n\n".join(text_parts)

            return full_text, structure_info

        except Exception as e:
            self.logger.error(f"Failed to extract content from Excel: {e}")
            raise

    def _extract_excel_insights(
        self, content: str, structure_info: dict, source_file: str
    ) -> list[Insight]:
        """Extract insights from Excel content."""
        insights = []

        # Structure insights
        if structure_info["sheet_names"]:
            insights.append(
                Insight(
                    content=f"Excel sheets: {', '.join(structure_info['sheet_names'])}",
                    source_file=source_file,
                    category="general",
                    confidence_score=0.8,
                    importance_score=0.5,
                )
            )

        # Look for numeric data patterns
        numbers = re.findall(r"\b\d+\.?\d*\b", content)
        if len(numbers) > 10:
            insights.append(
                Insight(
                    content=f"Contains numerical data ({len(numbers)} values)",
                    source_file=source_file,
                    category="general",
                    confidence_score=0.7,
                    importance_score=0.4,
                )
            )

        return insights


class CSVParser(FileParser):
    """Parser for CSV files."""

    def __init__(self):
        super().__init__()
        self.supported_extensions = {".csv", ".tsv"}
        self.supported_mime_types = {"text/csv", "text/tab-separated-values"}

    async def parse_file(self, file_path: Path) -> ResearchFile:
        """Parse CSV file and extract content."""
        try:
            metadata = self._get_file_metadata(file_path)

            # Read CSV content
            with open(file_path, encoding="utf-8", errors="ignore") as f:
                content = f.read(1024 * 100)  # Limit to 100KB

            # Extract basic structure
            lines = content.split("\n")[:50]  # First 50 lines

            # Detect delimiter
            delimiter = "," if "," in lines[0] else "\t" if "\t" in lines[0] else ";"

            # Process lines
            processed_lines = []
            for i, line in enumerate(lines):
                if line.strip():
                    if i == 0:  # Header
                        processed_lines.append(f"Headers: {line}")
                    elif i < 10:  # Sample data
                        processed_lines.append(line)
                    else:
                        break

            text_content = "\n".join(processed_lines)

            # Extract insights
            insights = self._extract_csv_insights(content, str(file_path))

            return ResearchFile(
                path=file_path,
                content=text_content,
                file_type="csv",
                metadata={
                    "original_format": "csv",
                    "delimiter": delimiter,
                    "estimated_rows": len(lines),
                    "estimated_columns": len(lines[0].split(delimiter)) if lines else 0,
                    **metadata.__dict__,
                },
                extracted_insights=[insight.content for insight in insights],
                confidence_score=0.5,
            )

        except Exception as e:
            self.logger.error(f"Failed to parse CSV file {file_path}: {e}")
            return ResearchFile(
                path=file_path,
                content="",
                file_type="csv",
                metadata={},
                processing_errors=[str(e)],
                confidence_score=0.0,
            )

    def _extract_csv_insights(self, content: str, source_file: str) -> list[Insight]:
        """Extract insights from CSV content."""
        insights = []

        lines = content.split("\n")
        if lines:
            header = lines[0]
            insights.append(
                Insight(
                    content=f"CSV data columns: {header}",
                    source_file=source_file,
                    category="general",
                    confidence_score=0.6,
                    importance_score=0.4,
                )
            )

        return insights


def get_advanced_parsers() -> list[FileParser]:
    """Get list of all available advanced parsers."""
    parsers = [
        CSVParser(),  # Always available
    ]

    if HAS_PYPDF2:
        parsers.append(PDFParser())

    if HAS_DOCX:
        parsers.append(DOCXParser())

    if HAS_OPENPYXL:
        parsers.append(ExcelParser())

    return parsers


def get_missing_dependencies() -> list[str]:
    """Get list of missing optional dependencies."""
    missing = []

    if not HAS_PYPDF2:
        missing.append("PyPDF2 (for PDF support)")

    if not HAS_DOCX:
        missing.append("python-docx (for DOCX support)")

    if not HAS_OPENPYXL:
        missing.append("openpyxl (for Excel support)")

    return missing
